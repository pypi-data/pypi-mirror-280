import pandas as pd
import numpy as np
import statsmodels.api as sm
from IPython.display import display, HTML
import importlib.resources as pkg_resources
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ._utils import ew_ret, vw_ret, remove_items, decimal, asterisk


def nw_ttest(df, ret, model_cols=[], maxlags=5):
    '''
    newey-west robust t test
    '''
    res_dict = {}
    
    dependent = df[ret]
    
    if model_cols:
        independent = sm.add_constant(df[model_cols])
    else:
        independent = pd.Series(np.ones_like(dependent), 
                                index=dependent.index, name="const")
    
    ols_model = sm.OLS(dependent, independent, missing="drop")
    reg_res = ols_model.fit(cov_type="HAC",
                            cov_kwds={'maxlags': maxlags}, kernel="bartlett")
    
    res_dict["mean"] = reg_res.params["const"]
    res_dict["tstats"] = reg_res.tvalues["const"]
    res_dict["pvalue"] = reg_res.pvalues["const"]
    
    return pd.Series(res_dict)

# sorting approach
class DoubleSorting(object):
    def __init__(self, data, sortbys, nqs, date, ret, mkt_cap):
        self.sortbys = sortbys
        self.nqs = nqs
        
        self.date = date
        self.ret = ret
        self.mkt_cap = mkt_cap
        
        columns = sortbys + [mkt_cap]
        self.data = data.dropna(subset=columns)
        self.data = self.data.reset_index(drop=True)
        
        self.data.index.name = "index"
        self.q_ret_name = "q_ret"
        
        
    def sorting(self, groupby, sortby, nq):
        # the smaller the label, the smaller the quantile
        labels = [i for i in range(1, nq+1)]
        
        groups = self.data.groupby(groupby, observed=False)
        quantiles = groups[sortby].apply(lambda x: pd.qcut(x, q=nq, labels=labels))
        
        quantiles = quantiles.reset_index(level=-1)
        quantiles = quantiles.reset_index(drop=True).set_index("index")
        return quantiles
    
    
    def quantile_return(self, groupby, vw):
        if vw:
            ret_func = vw_ret
        else:
            ret_func = ew_ret
            
        groups = self.data.groupby(groupby, observed=False)
        quantile_ret = groups.apply(lambda x: ret_func(data=x,
                                                       ret_col=self.ret,
                                                       mkt_cap_col=self.mkt_cap))
        quantile_ret.name = self.q_ret_name
        quantile_ret = quantile_ret.reset_index()
        return quantile_ret
    
    
    def sequential(self, vw=False):
        sortby1, sortby2 = self.sortbys
        name1, name2 = sortby1, sortby2
        nq1, nq2 = self.nqs
        
        groupby = [self.date]
        
        # sorting by variable 1
        self.data[name1] = self.sorting(groupby, sortby1, nq1)
        groupby.append(name1)
        # sorting by variable 2, sequentially
        self.data[name2] = self.sorting(groupby, sortby2, nq2)
        groupby.append(name2)
        
        # portfolio return for each sorting group
        quantile_ret = self.quantile_return(groupby, vw=vw)
        
        # wrap up results
        notes = [f"Sequential sorting: {sortby1} - {sortby2}"]
        results = SortingResults(quantile_ret, self.q_ret_name, 
                                 self.date, self.sortbys, notes)
        return results
    
    def independent(self, vw=False):
        sortby1, sortby2 = self.sortbys
        name1, name2 = sortby1, sortby2
        nq1, nq2 = self.nqs
        
        groupby = [self.date]
        
        # sorting by variable 1
        self.data[name1] = self.sorting(groupby, sortby1, nq1)
        # sorting by variable 2, independently
        self.data[name2] = self.sorting(groupby, sortby2, nq2)
        
        groupby.append(name1)
        groupby.append(name2)
        
        # portfolio return for each sorting group
        quantile_ret = self.quantile_return(groupby, vw=vw)
        
        # wrap up results
        notes = [f"Independent sorting: {sortby1}, {sortby2}"]
        results = SortingResults(quantile_ret, self.q_ret_name, 
                                 self.date, self.sortbys, notes)
        return results

    
# results wrapper
class SortingResults(object):
    def __init__(self, quantile_ret, qret, date, sortbys, init_notes=[]):
        self.res = quantile_ret
        self.notes = init_notes
        self.sortbys = sortbys
        self.qret, self.date = qret, date
        
        
    def inline_render(self, htmls):
        container = '<div style="display:flex;">'
        for n, html in enumerate(htmls):
            container += '<div style="margin-right: 20px;">'
            container += html
            container += '</div>'
        container += '</div>'
        return container
    
    
    def _make_table(self, rf_df, rf, alpha_models, layout, strategies, **kwargs):
        main = MainTable(self.res, self.qret, rf_df, rf, 
                         self.date, self.sortbys, layout=layout, **kwargs)
        
        if layout == "default":
            orientation = 'horizontal', "vertical"
        elif layout == "reverse":
            orientation = 'vertical', "horizontal"
            
        strategy1 = StrategyTable(self.res, self.qret, self.date, 
                                  self.sortbys[0], self.sortbys[1], alpha_models, 
                                  orientation[0], strategies[0], **kwargs)
        strategy2 = StrategyTable(self.res, self.qret, self.date, 
                                  self.sortbys[1], self.sortbys[0], alpha_models, 
                                  orientation[1], strategies[1], **kwargs)
        
        self.main = main
        self.strategies = [strategy1, strategy2]
        
        
    def _html_render(self, layout):
        main = self.main
        s1, s2 = self.strategies
        
        main_render = HtmlRenderer(main.means, main.tstats, main.hhead, 
                                   main.vhead, self.notes)
        main_html = main_render.render()
        
        render1 = HtmlRenderer(s1.means, s1.tstats, s1.hhead, 
                               s1.vhead, s1.notes)
        strategy1_html = render1.render()
        
        render2 = HtmlRenderer(s2.means, s2.tstats, s2.hhead, 
                               s2.vhead, s2.notes)
        strategy2_html = render2.render()
        
        if layout == "default":
            html1 = self.inline_render([main_html, strategy2_html])
            html2 = strategy1_html
        elif layout == "reverse":
            html1 = self.inline_render([main_html, strategy1_html])
            html2 = strategy2_html
            
        display(HTML(html1))
        display(HTML(html2))
        
        
    def _docx_save(self, output_path):
        main = self.main
        s1, s2 = self.strategies
        
        main_docx = DocxRenderer(main.means, main.tstats, main.hhead, main.vhead, self.notes)
        doc = main_docx.render()
        
        strategy1_docx = DocxRenderer(s1.means, s1.tstats, s1.hhead, 
                                      s1.vhead, s1.notes, doc=doc)
        doc = strategy1_docx.render()
        
        strategy2_docx = DocxRenderer(s2.means, s2.tstats, s2.hhead, 
                                      s2.vhead, s2.notes, doc=doc)
        doc = strategy2_docx.render()
        
        doc.save(output_path)
        print(f"\n\nFile saved at {output_path}")
        
        
    def summary(self, rf_df, rf, alpha_models, layout="default", 
                strategies=["HML", "HML"], output_path=None, **kwargs):
        '''
        -- rf_df: 
        a dataframe of risk-free rate and corresponding date
        date column should be consistent with the data in sorting approach
        
        -- rf: column of risk-free rate in rf_df
        
        -- alpha_models:
        a sequence of dataframes, each represents a pricing model used to calculate alpha
        date column should be consistent with the data in sorting approach
        
        -- strategies:
        a sequence of 'HML' or 'LMH', strategy for each sorting variable
        'HML' represents hign minus low, 'LMH' represents low minus high
        
        -- output_path:
        output results to a Microsoft Word file
        
        -- kwargs
        show_t: show tstats in main table
        show_stars: show asterisk in main table
        mean_decimal: decimal for means
        t_decimal: decimal for tstats
        '''
        
        self._make_table(rf_df, rf, alpha_models, layout, strategies, **kwargs)
        self._html_render(layout)
        if output_path:
            self._docx_save(output_path)
            
    
# make tables
class MainTable(object):
    def __init__(self, qret_df, qret, rf_df, rf, date, sortbys, layout="default", **kwargs):
        self.excess = "excess"
        self.maxlags = 5
        
        qret_df = self.add_excess(qret_df, qret, rf_df, rf, date)
        test_res = self.test_mean(qret_df, self.excess, sortbys, layout)
        self.means = self.mean_table(test_res, **kwargs)
        self.tstats = self.tstats_table(test_res, **kwargs)
        
    
    def add_excess(self, qret_df, qret, rf_df, rf, date):
        qret_df = qret_df.merge(rf_df, on=date, how="left")
        qret_df[self.excess] = qret_df[qret] - qret_df[rf]
        
        return qret_df
    
    
    def test_mean(self, qret_df, excess, sortbys, layout="default"):
        groups = qret_df.groupby(sortbys, observed=False)
        test_res = groups.apply(lambda x: nw_ttest(x, excess, [], maxlags=self.maxlags))
        
        if layout == "default":
            test_res = test_res.unstack(-1)
            self.vhead, self.hhead = sortbys
        elif layout == "reverse":
            test_res = test_res.unstack(0)
            self.hhead, self.vhead = sortbys
        else:
            raise ValueError("Valid layout parameter: 'default' or 'reverse'")
            
        return test_res
    
    
    def mean_table(self, test_res, show_stars=False, mean_decimal=3, **kwargs):
        means = test_res["mean"].map(lambda x:
                                     decimal(x * 100, mean_decimal) + "%")
        
        if show_stars:
            stars = test_res["pvalue"].map(asterisk)
            means += stars
            
        return means
    
    
    def tstats_table(self, test_res, show_t=False, t_decimal=3, **kwargs):
        if show_t:
            tstats = test_res["tstats"].map(lambda x:
                                            f"({decimal(x, t_decimal)})")
        else:
            tstats = test_res["tstats"].map(lambda x: "-")
            
        return tstats
    
    
class StrategyTable(object):
    def __init__(self, qret_df, qret, date, 
                 sortby_strategy, sortby_other, 
                 alpha_models,
                 orientation="vertical", strategy="HML", **kwargs):
        
        self.alpha_prefix = "alpha "
        self.excess = "excess"
        self.maxlags = 5
        self.notes = []
        
        # calculate the difference of return between high and low portfolios
        strategy_df = self.add_diff(qret_df, qret, date, 
                                    sortby_strategy, sortby_other, 
                                    strategy)
        
        alpha_models = [None] + alpha_models  # None for no model, i.e. excess return
        
        # iters models and caculate corresponding alphas
        self.get_alphas(date, strategy_df, sortby_other, alpha_models, **kwargs)
        # adjust according to orientation, add hhead and vhead
        self.orientation_adjust(orientation, sortby_strategy, sortby_other)
    
    def add_diff(self, qret_df, qret, date, sortby_strategy, sortby_other, strategy):
        qret_df = qret_df.set_index([date, sortby_other, sortby_strategy])
        qret_df = qret_df.unstack(sortby_strategy)[qret]
        
        columns = qret_df.columns
        hlabel, llabel = columns.max(), columns.min()
        
        if strategy == "HML":
            qret_df[self.excess] = qret_df[hlabel] - qret_df[llabel]
            self.notes.append(f"{sortby_strategy} strategy: high minus low")
        elif strategy == "LMH":
            qret_df[self.excess] = qret_df[llabel] - qret_df[hlabel]
            self.notes.append(f"{sortby_strategy} strategy: low minus high")
        else:
            raise ValueError("Valid strategy parameters: 'HML' or 'LMH'")
            
        strategy_df = qret_df[self.excess].reset_index()
        return strategy_df
    
    
    def test_mean(self, date, strategy_df, sortby_other, model=None):
        if isinstance(model, pd.DataFrame):
            model_cols = remove_items([date], model.columns)
            strategy_df = strategy_df.merge(model, on=date, how="left")
        else:
            model_cols = []
            strategy_df = strategy_df.copy()
            
        groups = strategy_df.groupby(sortby_other, observed=False)
        test_res = groups.apply(lambda x: nw_ttest(x, self.excess, model_cols, maxlags=self.maxlags))
        
        return test_res, model_cols
    
    
    def mean_table(self, test_res, show_stars_strategy=True, 
                   mean_decimal=3, **kwargs):
        means = test_res["mean"].map(lambda x:
                                     decimal(x * 100, mean_decimal) + "%")
        
        if show_stars_strategy:
            stars = test_res["pvalue"].map(asterisk)
            means += stars
            
        return means
    
    
    def tstats_table(self, test_res, show_t_strategy=True, 
                     t_decimal=3, **kwargs):
        if show_t_strategy:
            tstats = test_res["tstats"].map(lambda x:
                                            f"({decimal(x, t_decimal)})")
        else:
            tstats = test_res["tstats"].map(lambda x: "-")
            
        return tstats
    
    
    def get_alphas(self, date, strategy_df, sortby_other, alpha_models, **kwargs):
        # wrap up functions: test_mean, mean_table and tstats_table
        # iters alpha models and calculate alpha
        means_ls = []
        tstats_ls = []
        for n, model in enumerate(alpha_models):
            test_res, model_cols = self.test_mean(date, strategy_df, sortby_other, 
                                                  model=model)
            
            if n == 0:
                name = self.excess
            else:
                name = self.alpha_prefix + str(n)
                # add note to show variables in the model
                self.notes.append(name + f" model: {' ,'.join(model_cols)}")

            means = self.mean_table(test_res, **kwargs)
            tstats = self.tstats_table(test_res, **kwargs)
            means.name = name
            tstats.name = name
            
            means_ls.append(means)
            tstats_ls.append(tstats)
        
        self.means = pd.concat(means_ls, axis=1)
        self.tstats = pd.concat(tstats_ls, axis=1)
        
    def orientation_adjust(self, orientation, sortby_strategy, sortby_other):
        if orientation == "vertical":
            self.hhead = f"{sortby_strategy} strategy"
            self.vhead = f"{sortby_other}"
        elif orientation == "horizontal":
            self.means = self.means.T
            self.tstats = self.tstats.T
            
            self.hhead = f"{sortby_other}"
            self.vhead = f"{sortby_strategy} strategy"
            
            
# table render
class HtmlRenderer(object):
    def __init__(self, means, tstats, hhead, vhead, notes=[]):
        self.means, self.tstats  = means, tstats
        self.hhead, self.vhead = hhead, vhead
        self.notes = notes
        
        self.hlabel, self.vlabel = means.columns, means.index
        
        self.table = '<table><tbody>'
        
        
    def cell(self, content="", bold=False):
        if bold:
            node = "<th>{}</th>"
        else:
            node = "<td>{}</td>"

        return node.format(content)
    
    def add_cells(self, series, bold=False, center=True):
        html = ""
        for value in series:
            html += self.cell(value, bold=bold)
        return html
    
    def _add_hhead(self):
        hhead = '<tr>' + self.cell() * 2  # skip two columns for vhead & vlabel
        hhead += f'<th colspan="{len(self.hlabel)}" style="text-align: center;">'  # hhead merge cells
        hhead += f'{self.hhead}</th></tr>'  # add hhead
        
        self.table += hhead
        
    def _add_hlabel(self):
        hlabel = '<tr>' + self.cell() * 2  # skip two columns for vhead & vlabel
        hlabel += self.add_cells(self.hlabel, bold=True) + '</tr>'  # add hlabel
        
        self.table += hlabel
        
    def _add_tbody(self):
        # iters rows and insert into html table
        for i in range(len(self.vlabel)):
            mrow = self.means.iloc[i]
            trow = self.tstats.iloc[i]

            if i == 0:  # first row add vhead
                # mean
                tbody = f'<tr><th rowspan="{len(self.vlabel) * 2}">'  # vhead merge cells
                tbody += f'{self.vhead}</th>'  # add vhead
                tbody += f'<th>{self.vlabel[i]}</th>'  # add vlabel
                tbody += self.add_cells(mrow) + '</tr>'  # add content
            else:
                # mean
                tbody += f"<tr><th>{self.vlabel[i]}</th>"  # add vlabel
                tbody += self.add_cells(mrow) + '</tr>'  # add content

            # tstats
            tbody += f"<tr><th></th>"  # skip vlabel cell
            tbody += self.add_cells(trow) + '</tr>'  # add content
        
        tbody += '</tbody>'  # close tbody tag
        self.table += tbody
            
    def add_notes(self):
        self.table += '<p style="font-size:12px; line-height:10px;">'
        self.table += 'Notes:'
        for i, note in enumerate(self.notes):
            self.table += '<p style="font-size:12px; line-height:4px;">'
            self.table += f'({i+1}) {note}</p>'
        
    def render(self):
        self._add_hhead()
        self._add_hlabel()
        self._add_tbody()
        
        self.table += '</table>'  # close table tag
        self.add_notes()
        return self.table
    
    
class DocxRenderer(object):
    def __init__(self, means, tstats, hhead, vhead, notes=[], doc=None, template="template.docx", style="academic table"):
        self.means, self.tstats  = means, tstats
        self.hhead, self.vhead = hhead, vhead
        self.notes = notes
        self.style = style
        
        self.hlabel, self.vlabel = means.columns, means.index
        
        if doc:
            self.doc = doc
        else:
            with pkg_resources.open_binary('empiricalfin', 'data/template.docx') as file:
                self.doc = Document(file)
            
        self.table = self.doc.add_table(rows=2, cols=len(self.hlabel)+2)
        
    @property
    def rows(self):
        return self.table.rows
    @property
    def cols(self):
        return self.table.columns
    
    def add_cells(self, series, cells):
        for text, cell in zip(series, cells):
            cell.text = text
    
    def _add_hhead(self):
        # merge cells for hhead
        row = self.rows[0]
        begin = row.cells[2]
        end = row.cells[-1]
        begin.merge(end)
        # add hhead
        begin.text = self.hhead
        
    def _add_hlabel(self):
        row = self.rows[1].cells[2:]
        for cell, text in zip(row, self.hlabel):
            cell.text = str(text)
            
    def _add_tbody(self):
        # each iteration add two rows, one for mean and another for tstats
        for i in range(len(self.vlabel)):
            mrow = self.means.iloc[i]
            trow = self.tstats.iloc[i]
            
            mcells = self.table.add_row().cells  # add row
            mcells[1].text = str(self.vlabel[i])  # add vlabel
            
            # add mean
            mcells = mcells[2:]
            self.add_cells(mrow, mcells)
            
            # add tstats
            tcells = self.table.add_row().cells[2:]
            self.add_cells(trow, tcells)
            
    def _add_vhead(self):
        col = self.cols[0]
        # merge cells for vhead
        begin = col.cells[2]
        end = col.cells[-1]
        begin.merge(end)
        # add vhead
        begin.text = self.vhead
        
    def _alignment(self):
        for row in self.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
    def _add_note(self):
        self.doc.add_paragraph("Notes:")
        for i, note in enumerate(self.notes):
            note = f"({i+1})" + note
            self.doc.add_paragraph(note)
        self.doc.add_paragraph("="*59)
        self.doc.add_paragraph(" ")
                    
    def render(self):
        self._add_hhead()
        self._add_hlabel()
        self._add_tbody()
        self._add_vhead()
        self._alignment()
        self.table.style = self.style
        self._add_note()
        return self.doc