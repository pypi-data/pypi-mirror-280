from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, Image, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch
from PIL import Image as PILImage

def convert_to_pdf(word_file):
    # Load the Word document
    doc = Document(word_file)
    
    # Create a buffer for PDF
    pdf_buffer = BytesIO()
    
    # Initialize the PDF content
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles for paragraphs
    custom_normal = ParagraphStyle(name='CustomNormal', parent=styles['Normal'], leading=12)
    custom_bold = ParagraphStyle(name='CustomBold', parent=styles['Normal'], leading=12, spaceAfter=10, fontName='Helvetica-Bold')
    custom_italic = ParagraphStyle(name='CustomItalic', parent=styles['Normal'], leading=12, spaceAfter=10, fontName='Helvetica-Oblique')
    custom_heading1 = ParagraphStyle(name='CustomHeading1', parent=styles['Heading1'], leading=14, spaceAfter=10, fontName='Helvetica-Bold')
    custom_heading2 = ParagraphStyle(name='CustomHeading2', parent=styles['Heading2'], leading=12, spaceAfter=10, fontName='Helvetica-Bold')
    
    # Loop through each paragraph in the Word document
    for paragraph in doc.paragraphs:
        # Extract the paragraph text and style
        text = ""
        for run in paragraph.runs:
            if run.bold:
                text += f'<b>{run.text}</b>'
            elif run.italic:
                text += f'<i>{run.text}</i>'
            else:
                text += run.text.replace('\n', '<br/>')
        
        # Create a ReportLab Paragraph object with appropriate style
        style_name = paragraph.style.name
        if style_name == 'Normal':
            style = custom_normal
        elif style_name == 'Heading 1':
            style = custom_heading1
        elif style_name == 'Heading 2':
            style = custom_heading2
        else:
            style = custom_normal
        
        elements.append(Paragraph(text, style))
        
        # Handle bullet points
        if paragraph.style.name == 'List Bullet':
            bullet_list = ListFlowable([ListItem(Paragraph(text, custom_normal))], bulletType='bullet')
            elements.append(bullet_list)
        
        # Add spacing after paragraphs
        elements.append(Spacer(1, 12))
    
    # Loop through each table in the Word document
    for table in doc.tables:
        # Extract data from table
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(Paragraph(cell.text, custom_normal))
            table_data.append(row_data)
        
        # Convert table data to PDF Table object
        pdf_table = Table(table_data)
        elements.append(pdf_table)
        elements.append(Spacer(1, 12))
    
    # Loop through each element in the document to find images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img = rel.target_part.blob
            img_buffer = BytesIO(img)
            pil_image = PILImage.open(img_buffer)
            img_width, img_height = pil_image.size
            aspect_ratio = img_width / img_height
            max_width, max_height = 4 * inch, 4 * inch  # Maximum image size
            if img_width > max_width or img_height > max_height:
                if img_width > img_height:
                    img_width = max_width
                    img_height = max_width / aspect_ratio
                else:
                    img_height = max_height
                    img_width = max_height * aspect_ratio
            img_element = Image(img_buffer, width=img_width, height=img_height)
            img_element.hAlign = 'CENTER'
            elements.append(img_element)
            elements.append(Spacer(1, 12))

    # Build the PDF document using SimpleDocTemplate
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    doc.build(elements)
    
    # Save PDF buffer
    pdf_buffer.seek(0)
    
    return pdf_buffer
