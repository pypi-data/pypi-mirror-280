#! /usr/bin/python3

# Brief WordX main program
# Author: Ashad Mohamed (aka. mashad)
# **********************************
# This the main file for this project, contain all general function
# to implement docx manipulation.

import json
import docx2python
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# TO-DO: Function Convert Docx To Json
# Convert DOCX to JSON
def convertDocxToJson(file_path: str, output_json_path: str) -> None:
    content = docx2python.docx2python(file_path)
    doc_dict = {
        'text': content.text,
        'comments': content.comments
    }
    with open(output_json_path, 'w') as json_file:
        json.dump(doc_dict, json_file, indent=4)

# TO-DO: Function Convert Json to Docx
def convertJsonToDocx(input_json_path: str, output_docx_path: str) -> None:
    with open(input_json_path, 'r') as json_file:
        doc_dict = json.load(json_file)
    doc = Document()
    doc.add_paragraph(doc_dict['text'])
    for comment in doc_dict['comments']:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(comment[0])
        comment_element = OxmlElement('w:comment')
        comment_element.set(qn('w:author'), comment[1])
        comment_element.set(qn('w:date'), comment[2])
        comment_element.text = comment[3]
        run._r.append(comment_element)
    doc.save(output_docx_path)
    
    
# TO-DO: Function Get All Comments
def getComments(file_path: str) -> []:
    comments = []
    with docx2python.docx2python(file_path) as docx_content:
        for comment in docx_content.comments:
            comments.append(comment)

        # text, author, timestamp, comment = comment
        # Print or process each component as needed
        # print(f"Text: {text.strip()}")
        # print(f"Author: {author}")
        # print(f"Timestamp: {timestamp}")
        # print(f"Comment: {comment.strip()}\n")

    comments_json = [
        {
            "text": comment[0].strip(),
            "author": comment[1],
            "timestamp": comment[2],
            "comment": comment[3].strip()
        }
        for comment in comments
    ]
    return json.dumps(comments_json, indent=4)

# TO-DO: Function Add Comment
# Add a comment
def addComment(file_path: str, text_to_comment: str, comment_text: str, author: str, timestamp: str) -> None:
    doc = Document(file_path)
    comment_id = 0

    # Create comments part if it does not exist
    comments_part = doc.part._element.xpath('/w:document/w:comments')
    if not comments_part:
        comments_part = OxmlElement('w:comments')
        doc.part._element.append(comments_part)
    else:
        comments_part = comments_part[0]

    new_comment = OxmlElement('w:comment')
    new_comment.set(qn('w:id'), str(comment_id))
    new_comment.set(qn('w:author'), author)
    new_comment.set(qn('w:date'), timestamp)
    new_comment.append(OxmlElement('w:p'))
    new_comment[0].append(OxmlElement('w:r'))
    new_comment[0][0].append(OxmlElement('w:t'))
    new_comment[0][0][0].text = comment_text
    comments_part.append(new_comment)

    for paragraph in doc.paragraphs:
        if text_to_comment in paragraph.text:
            start = paragraph.text.find(text_to_comment)
            end = start + len(text_to_comment)
            run = paragraph.add_run()
            comment_range_start = OxmlElement('w:commentRangeStart')
            comment_range_start.set(qn('w:id'), str(comment_id))
            run._r.insert(0, comment_range_start)
            run = paragraph.add_run()
            comment_range_end = OxmlElement('w:commentRangeEnd')
            comment_range_end.set(qn('w:id'), str(comment_id))
            run._r.append(comment_range_end)
            run = paragraph.add_run()
            comment_reference = OxmlElement('w:commentReference')
            comment_reference.set(qn('w:id'), str(comment_id))
            run._r.append(comment_reference)

    doc.save(file_path)


# TO-DO: Function Tracking Modifications
def trackModifications(file_path: str) -> dict:
return {}