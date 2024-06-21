from docx import Document
from docx.shared import RGBColor, Pt, Cm
import os
import glob
import docx ,TimeUtil
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Inches
import LogUtil,util.StringUtil as StringUtil

# path="word/src/src.docx"

mkdir_path = "word/src"
if not os.path.exists(mkdir_path):
    os.mkdir(mkdir_path)

keyword = 'CATD'
file_path = "word/src.docx"

docx = Document(file_path)

print(docx)

# 4-左右两边的标尺分别对应“6”和“50”
# docx.sections[1].left_margin = Cm(6)
# docx.sections[1].right_margin = Cm(6)

def addTable(table):
    print("行",len(table.rows))
    print("列",len(table.columns))
    csv=Document()
    tb = csv.add_table(rows=len(table.rows), cols=len(table.columns))
    tb.add_row()
    for row in range(len(table.rows)):
        for col in range(len(table.columns)):
            tb.cell(row,col).width=1
            tb.cell(row,col).text=table.cell(row,col).text
            tb.cell(row,col).width=Cm(6)
    # tb.style='Light Shading Accent 2'
    tb.style='Table Grid'
    tb.autofit=True
    tableName=StringUtil.commonStrip(table.cell(0,0).text)
    csv.save("word/dest/"+TimeUtil.getFoarmatNanoTimestamp()+"-"+tableName+'.docx')
tables = docx.tables
for i in range(len(tables)):
    table = tables[i]
    addTable(table)
# addTable(tables[1])



